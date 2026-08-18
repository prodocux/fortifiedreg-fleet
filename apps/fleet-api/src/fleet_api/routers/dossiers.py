"""
Dossier Management Router.
Handles dossier creation, document content registration, 5-format binary profiling,
SCCS toxicology evaluation, plan compilation, and execution orchestration.
"""
import base64
import csv
import hashlib
import io
import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID
from fastapi import APIRouter, Depends, HTTPException, status
from pydantic import BaseModel, Field

from fleet_api.deps import (
    get_audit_log,
    get_checkpoint_store,
    get_document_resolver,
    get_orchestrator,
    get_resume_context_store,
    get_tenant_and_actor,
)
from fleet_domain_cosmetics.inci_verifier import evaluate_inci_compliance
from fleet_domain_cosmetics.mos_calculator import (
    calculate_mos,
    calculate_sed,
    evaluate_toxicology_mos,
)
from fleet_governance_core.models.approval import (
    AuthenticatedActor,
    PDXApprovalRequest,
    PDXWorkflowCheckpoint,
)
from fleet_governance_core.models.audit import AuditEvent, AuditEventTypeEnum
from fleet_governance_core.models.case import DossierCase
from fleet_governance_core.models.hashing import compute_data_sha256
from fleet_governance_core.models.verifier import VerifierStatusEnum
from fleet_governance_core.ports.audit_log_port import AuditLogPort
from fleet_governance_core.ports.checkpoint_store_port import CheckpointStorePort
from fleet_governance_core.ports.document_resolver_port import DocumentResolverPort
from fleet_governance_core.ports.orchestrator_port import ExecutionOrchestratorPort
from fleet_governance_core.ports.resume_context_store_port import ResumeContextStorePort

router = APIRouter(prefix="/v1/dossiers", tags=["Dossiers"])

CASES_DB: Dict[str, Dict[str, DossierCase]] = {}  # tenant_id -> {case_id_str: case}

MAX_BASE64_CHARS = 68_000_000  # ~50 MiB raw ceiling
FORMAT_LIMITS = {
    ".pdf": 10 * 1024 * 1024,
    ".docx": 16 * 1024 * 1024,
    ".csv": 8 * 1024 * 1024,
    ".xlsx": 16 * 1024 * 1024,
    ".pptx": 32 * 1024 * 1024,
}


class DocumentRegistrationRequest(BaseModel):
    doc_id: str = Field(min_length=1, max_length=64, pattern=r"^[a-zA-Z0-9_-]+$")
    content_b64: str = Field(min_length=1, max_length=MAX_BASE64_CHARS)
    filename: Optional[str] = Field(default=None, min_length=1, max_length=255, pattern=r"^[a-zA-Z0-9_.-]+$")
    expected_sha256: Optional[str] = Field(default=None, pattern=r"^[a-fA-F0-9]{64}$")


class DocumentProfileRequest(BaseModel):
    doc_id: str = Field(min_length=1, max_length=64, pattern=r"^[a-zA-Z0-9_-]+$")
    filename: str = Field(min_length=1, max_length=255, pattern=r"^[a-zA-Z0-9_.-]+$")
    content_b64: str = Field(min_length=1, max_length=MAX_BASE64_CHARS)


@router.post("/documents/register", response_model=Dict[str, Any])
def register_document(
    req: DocumentRegistrationRequest,
    identity: Tuple[str, AuthenticatedActor] = Depends(get_tenant_and_actor),
    resolver: DocumentResolverPort = Depends(get_document_resolver),
) -> Dict[str, Any]:
    """Register raw binary document content for a supplier document in the tenant's resolver."""
    tenant_id, actor = identity

    if not req.doc_id or not req.doc_id.strip():
        raise HTTPException(status_code=400, detail="Invalid doc_id provided.")

    if not req.content_b64 or len(req.content_b64) > MAX_BASE64_CHARS:
        raise HTTPException(status_code=400, detail="Document payload exceeds maximum allowed request ceiling.")

    try:
        raw_bytes = base64.b64decode(req.content_b64, validate=True)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Base64 payload encoding.")

    if len(raw_bytes) == 0:
        raise HTTPException(status_code=400, detail="Document content cannot be empty.")

    filename = req.filename or f"{req.doc_id}.pdf"
    ext = Path(filename).suffix.casefold()
    if ext not in FORMAT_LIMITS:
        raise HTTPException(status_code=400, detail="Unsupported document format.")

    limit = FORMAT_LIMITS[ext]
    if len(raw_bytes) > limit:
        raise HTTPException(status_code=400, detail="Document payload exceeds maximum allowed size for format.")

    try:
        sha256_digest = resolver.register_document(
            tenant_id=tenant_id,
            doc_id=req.doc_id,
            content=raw_bytes,
            filename=filename,
            expected_sha256=req.expected_sha256,
        )
    except ValueError:
        raise HTTPException(status_code=409, detail="Document already registered with different content.")

    return {
        "status": "registered",
        "doc_id": req.doc_id,
        "tenant_id": tenant_id,
        "filename": filename,
        "sha256": sha256_digest,
        "size_bytes": len(raw_bytes),
    }


@router.post("/documents/profile", response_model=Dict[str, Any])
def profile_document(
    req: DocumentProfileRequest,
    identity: Tuple[str, AuthenticatedActor] = Depends(get_tenant_and_actor),
) -> Dict[str, Any]:
    """Parse real binary document content and extract structural metadata properties."""
    tenant_id, _ = identity
    ext = Path(req.filename).suffix.casefold()

    if ext not in FORMAT_LIMITS:
        raise HTTPException(status_code=400, detail=f"Unsupported format '{ext}'. Must be PDF, DOCX, CSV, XLSX, or PPTX.")

    try:
        raw_bytes = base64.b64decode(req.content_b64, validate=True)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Base64 payload.")

    if len(raw_bytes) == 0:
        raise HTTPException(status_code=400, detail="Document content cannot be empty.")

    if len(raw_bytes) > FORMAT_LIMITS[ext]:
        raise HTTPException(status_code=400, detail=f"Document exceeds maximum limit of {FORMAT_LIMITS[ext]} bytes.")

    sha256_raw = hashlib.sha256(raw_bytes).hexdigest()
    profile_data: Dict[str, Any] = {
        "doc_id": req.doc_id,
        "filename": req.filename,
        "format": ext.lstrip(".").upper(),
        "size_bytes": len(raw_bytes),
        "raw_sha256": sha256_raw,
    }

    # Validate Magic Bytes & Extract Real Structural Profile
    try:
        if ext == ".pdf":
            if not raw_bytes.startswith(b"%PDF"):
                raise HTTPException(status_code=400, detail="Magic bytes mismatch: Not a valid PDF file.")
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            profile_data["page_count"] = len(reader.pages)
            profile_data["is_encrypted"] = reader.is_encrypted

        elif ext == ".docx":
            if not raw_bytes.startswith(b"PK\x03\x04"):
                raise HTTPException(status_code=400, detail="Magic bytes mismatch: Not a valid DOCX container.")
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            profile_data["paragraph_count"] = len(doc.paragraphs)
            profile_data["table_count"] = len(doc.tables)

        elif ext == ".xlsx":
            if not raw_bytes.startswith(b"PK\x03\x04"):
                raise HTTPException(status_code=400, detail="Magic bytes mismatch: Not a valid XLSX container.")
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True)
            profile_data["sheet_names"] = wb.sheetnames
            profile_data["sheet_count"] = len(wb.sheetnames)

        elif ext == ".pptx":
            if not raw_bytes.startswith(b"PK\x03\x04"):
                raise HTTPException(status_code=400, detail="Magic bytes mismatch: Not a valid PPTX container.")
            import pptx
            prs = pptx.Presentation(io.BytesIO(raw_bytes))
            profile_data["slide_count"] = len(prs.slides)

        elif ext == ".csv":
            text = raw_bytes.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
            profile_data["row_count"] = len(rows)
            profile_data["column_count"] = len(rows[0]) if rows else 0
            profile_data["header_columns"] = rows[0] if rows else []

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse {ext.upper()} structure: {str(exc)}",
        )

    profile_digest = hashlib.sha256(json.dumps(profile_data, sort_keys=True).encode("utf-8")).hexdigest()
    profile_data["profile_digest"] = profile_digest
    return profile_data


@router.post("/evaluate-sccs", response_model=Dict[str, Any])
def evaluate_sccs_compliance(
    case: DossierCase,
    identity: Tuple[str, AuthenticatedActor] = Depends(get_tenant_and_actor),
) -> Dict[str, Any]:
    """Run real server-side SCCS 12th Notes of Guidance toxicology and Annex II/V compliance evaluation."""
    tenant_id, _ = identity
    if case.tenant_id != tenant_id:
        raise HTTPException(
            status_code=403,
            detail=f"Payload tenant_id '{case.tenant_id}' does not match authenticated token tenant '{tenant_id}'",
        )

    inci_res = evaluate_inci_compliance(case)
    mos_res = evaluate_toxicology_mos(case)

    # Determine overall verifier status
    if inci_res.status == VerifierStatusEnum.FAIL or mos_res.status == VerifierStatusEnum.FAIL:
        overall_status = "fail"
    elif mos_res.status == VerifierStatusEnum.REVIEW or inci_res.status == VerifierStatusEnum.REVIEW:
        overall_status = "review"
    else:
        overall_status = "pass"

    exp = case.exposure_scenario
    substance_evaluations: List[Dict[str, Any]] = []
    for item in case.formula:
        if item.inci_name.upper() == "AQUA":
            substance_evaluations.append({
                "inci_name": item.inci_name,
                "concentration_pct": item.concentration_pct,
                "cas_number": item.cas_number,
                "noael_mg_kg_day": item.noael_mg_kg_day,
                "sed_mg_kg_day": 0.0,
                "margin_of_safety": None,
                "status": "exempt",
            })
            continue

        sed = calculate_sed(
            daily_applied_amount_g=exp.daily_applied_amount_g,
            concentration_pct=item.concentration_pct,
            retention_factor=exp.retention_factor,
            body_weight_kg=exp.body_weight_kg,
        )
        mos = calculate_mos(item.noael_mg_kg_day, sed) if (item.noael_mg_kg_day is not None and item.noael_mg_kg_day > 0) else None

        if mos is None:
            status_item = "review"
        elif mos < 100.0:
            status_item = "fail"
        else:
            status_item = "pass"

        substance_evaluations.append({
            "inci_name": item.inci_name,
            "concentration_pct": item.concentration_pct,
            "cas_number": item.cas_number,
            "noael_mg_kg_day": item.noael_mg_kg_day,
            "sed_mg_kg_day": sed,
            "margin_of_safety": mos,
            "status": status_item,
        })

    evidence_summary = {
        "case_id": str(case.case_id),
        "product_name": case.product_name,
        "overall_status": overall_status,
        "inci_result": inci_res.model_dump(mode="json"),
        "mos_result": mos_res.model_dump(mode="json"),
        "substances": substance_evaluations,
    }
    evidence_digest = hashlib.sha256(json.dumps(evidence_summary, sort_keys=True).encode("utf-8")).hexdigest()

    return {
        "case_id": str(case.case_id),
        "product_name": case.product_name,
        "rule_set_version": "SCCS_12TH_NOTES_OF_GUIDANCE_2025.1",
        "verifier_status": overall_status,
        "inci_compliance": inci_res.model_dump(mode="json"),
        "toxicology_mos": mos_res.model_dump(mode="json"),
        "substance_evaluations": substance_evaluations,
        "exposure_summary": {
            "daily_applied_amount_g": exp.daily_applied_amount_g,
            "retention_factor": exp.retention_factor,
            "body_weight_kg": exp.body_weight_kg,
        },
        "evidence_digest": evidence_digest,
    }


@router.post("/create", response_model=Dict[str, Any])
def create_dossier(
    case: DossierCase,
    identity: Tuple[str, AuthenticatedActor] = Depends(get_tenant_and_actor),
    audit_log: AuditLogPort = Depends(get_audit_log),
    resume_context_store: ResumeContextStorePort = Depends(get_resume_context_store),
) -> Dict[str, Any]:
    """Create a new product regulatory dossier case and compute canonical digest."""
    tenant_id, actor = identity

    if case.tenant_id != tenant_id:
        raise HTTPException(
            status_code=403,
            detail=f"Payload tenant_id '{case.tenant_id}' does not match authenticated token tenant '{tenant_id}'",
        )

    case_id_str = str(case.case_id)
    CASES_DB.setdefault(tenant_id, {})[case_id_str] = case
    if hasattr(resume_context_store, "save_case"):
        resume_context_store.save_case(tenant_id, case)
    case_digest = compute_data_sha256(case)

    audit_log.append_audit_event(
        AuditEvent(
            tenant_id=tenant_id,
            run_id=f"run-pif-{case_id_str}",
            event_type=AuditEventTypeEnum.CASE_CREATED,
            actor_id=actor.sub,
            payload={
                "case_id": case_id_str,
                "product_name": case.product_name,
                "case_digest": case_digest,
            },
        )
    )

    return {
        "status": "created",
        "case_id": case_id_str,
        "product_name": case.product_name,
        "case_digest": case_digest,
    }


@router.post("/{case_id}/compile-and-run", response_model=Dict[str, Any])
def compile_and_run_dossier(
    case_id: UUID,
    identity: Tuple[str, AuthenticatedActor] = Depends(get_tenant_and_actor),
    orch: ExecutionOrchestratorPort = Depends(get_orchestrator),
    checkpoint_store: CheckpointStorePort = Depends(get_checkpoint_store),
    audit_log: AuditLogPort = Depends(get_audit_log),
    resume_context_store: ResumeContextStorePort = Depends(get_resume_context_store),
) -> Dict[str, Any]:
    """Compile dossier case into PDX plan and run through deterministic verifiers."""
    tenant_id, actor = identity
    case_id_str = str(case_id)
    case = CASES_DB.get(tenant_id, {}).get(case_id_str)
    if not case and hasattr(resume_context_store, "get_case"):
        case = resume_context_store.get_case(tenant_id, case_id_str)
    if not case:
        raise HTTPException(status_code=404, detail=f"Dossier case '{case_id_str}' not found.")

    case_payload = case.model_dump(mode="json")
    plan = orch.compile_execution_plan(case_payload)
    plan_digest = compute_data_sha256(plan)

    audit_log.append_audit_event(
        AuditEvent(
            tenant_id=tenant_id,
            run_id=plan.get("request_id", f"run-{case_id_str[:8]}"),
            event_type=AuditEventTypeEnum.PLAN_COMPILED,
            actor_id=actor.sub,
            payload={"plan_digest": plan_digest},
        )
    )

    exec_result = orch.execute_plan(plan, case_payload=case_payload)

    if exec_result.get("status") == "awaiting_approval":
        chk_dict = exec_result["checkpoint"]
        checkpoint = PDXWorkflowCheckpoint.model_validate(chk_dict)
        checkpoint_store.save_checkpoint(tenant_id, checkpoint)

        if "approval_request" in exec_result:
            req_dict = exec_result["approval_request"]
            approval_req = PDXApprovalRequest.model_validate(req_dict)
            checkpoint_store.save_approval_request(tenant_id, approval_req)

        audit_log.append_audit_event(
            AuditEvent(
                tenant_id=tenant_id,
                run_id=checkpoint.run_id,
                event_type=AuditEventTypeEnum.CHECKPOINT_CREATED,
                actor_id=actor.sub,
                payload={
                    "checkpoint_id": checkpoint.checkpoint_id,
                    "approval_request_id": exec_result.get("approval_request_id"),
                },
            )
        )

    return {
        "case_id": case_id_str,
        "plan": plan,
        "plan_digest": plan_digest,
        "execution": exec_result,
    }


@router.get("/{case_id}", response_model=Dict[str, Any])
def get_dossier(
    case_id: UUID,
    identity: Tuple[str, AuthenticatedActor] = Depends(get_tenant_and_actor),
    resume_context_store: ResumeContextStorePort = Depends(get_resume_context_store),
) -> Dict[str, Any]:
    """Retrieve dossier case details and canonical digest."""
    tenant_id, _ = identity
    case_id_str = str(case_id)
    case = CASES_DB.get(tenant_id, {}).get(case_id_str)
    if not case and hasattr(resume_context_store, "get_case"):
        case = resume_context_store.get_case(tenant_id, case_id_str)
    if not case:
        raise HTTPException(status_code=404, detail=f"Dossier case '{case_id_str}' not found.")

    return {
        "case": case.model_dump(mode="json"),
        "case_digest": compute_data_sha256(case),
    }
